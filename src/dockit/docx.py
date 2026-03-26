"""Word document processing — text formatting, extraction, style cleanup.

Pure logic module: bytes in, bytes out. No file paths, no console output.

Usage:
    from dockit.docx import format_text, extract_text, cleanup_styles

    result = format_text(doc_bytes)        # Fix quotes/punctuation/units
    md = extract_text(doc_bytes)           # Extract as Markdown
    result = cleanup_styles(doc_bytes)     # Remove unused styles
"""

import copy
import re
import zipfile
from dataclasses import dataclass, field
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from dockit.text import fix_all


@dataclass
class FormatResult:
    """Result of a document formatting operation."""

    data: bytes
    stats: dict[str, int] = field(default_factory=dict)


# -- Internal helpers ----------------------------------------------------------

QUOTE_CHARS = {"\u201c", "\u201d"}


def _set_run_font(run_element, font_name: str):
    """Set font for a run element (ascii + hAnsi + eastAsia)."""
    rPr = run_element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run_element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:hint"), "eastAsia")


def _split_run_at_quotes(run):
    """Split a run at quote characters, returning [(text, is_quote), ...] or None."""
    text = run.text
    if not text or not any(c in QUOTE_CHARS for c in text):
        return None

    segments = []
    buf = []
    for c in text:
        if c in QUOTE_CHARS:
            if buf:
                segments.append(("".join(buf), False))
                buf = []
            segments.append((c, True))
        else:
            buf.append(c)
    if buf:
        segments.append(("".join(buf), False))
    return segments


def _apply_quote_split(run, segments, quote_font: str):
    """Split run into multiple runs: quote chars get the specified font."""
    parent = run._element.getparent()

    # First segment reuses the original run
    first_text, first_is_quote = segments[0]
    run.text = first_text
    if first_is_quote:
        _set_run_font(run._element, quote_font)

    # Subsequent segments: deep-copy original run, insert after
    insert_after = run._element
    for seg_text, is_quote in segments[1:]:
        new_r = copy.deepcopy(run._element)
        # Clear old text elements from the copy
        for t_elem in new_r.findall(qn("w:t")):
            new_r.remove(t_elem)
        t_elem = OxmlElement("w:t")
        t_elem.text = seg_text
        t_elem.set(qn("xml:space"), "preserve")
        new_r.append(t_elem)

        if is_quote:
            _set_run_font(new_r, quote_font)
        else:
            # Non-quote segments: restore original font
            rPr = new_r.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                orig_rPr = run._element.find(qn("w:rPr"))
                orig_rFonts = orig_rPr.find(qn("w:rFonts")) if orig_rPr is not None else None
                if rFonts is not None and orig_rFonts is not None:
                    rPr.replace(rFonts, copy.deepcopy(orig_rFonts))
                elif rFonts is not None and orig_rFonts is None:
                    rPr.remove(rFonts)

        parent.insert(list(parent).index(insert_after) + 1, new_r)
        insert_after = new_r


def _process_paragraph(paragraph, stats: dict, quote_counter: int, quote_font: str) -> int:
    """Process a single paragraph: fix text and split quotes into separate runs.

    Returns updated quote_counter.
    """
    if not paragraph.runs:
        return quote_counter

    original_runs = list(paragraph.runs)

    for run in original_runs:
        if not run.text:
            continue

        original_text = run.text
        fixed_text, fix_stats, quote_counter = fix_all(original_text, quote_counter)

        stats["quotes"] += fix_stats["quotes"]
        stats["punctuation"] += fix_stats["punctuation"]
        stats["units"] += fix_stats["units"]

        if fixed_text != original_text:
            run.text = fixed_text

        # Split quotes into separate runs with specified font
        segments = _split_run_at_quotes(run)
        if segments:
            _apply_quote_split(run, segments, quote_font)

    return quote_counter


def _process_table(table, stats: dict, quote_counter: int, quote_font: str) -> int:
    """Process all paragraphs in a table. Returns updated quote_counter."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                quote_counter = _process_paragraph(paragraph, stats, quote_counter, quote_font)
    return quote_counter


# -- Public API ----------------------------------------------------------------


def format_text(
    doc_bytes: bytes,
    *,
    fix_quotes: bool = True,
    fix_punctuation: bool = True,
    fix_units: bool = True,
    quote_font: str = "\u5b8b\u4f53",
    process_headers_footers: bool = True,
    strip_headers_footers: bool = False,
) -> FormatResult:
    """Format text in a Word document.

    Fixes quotes (with font splitting), punctuation, and unit symbols
    across all paragraphs, tables, headers, and footers.

    Args:
        doc_bytes: Raw bytes of a .docx file.
        fix_quotes: Whether to fix quote characters.
        fix_punctuation: Whether to convert English punctuation to Chinese.
        fix_units: Whether to convert Chinese unit names to symbols.
        quote_font: Font name for quote characters (default: Song Ti).
        process_headers_footers: Whether to process headers/footers.
        strip_headers_footers: If True, remove all headers/footers entirely.

    Returns:
        FormatResult with processed document bytes and replacement stats.
    """
    doc = Document(BytesIO(doc_bytes))
    stats = {"quotes": 0, "punctuation": 0, "units": 0}
    quote_counter = 0

    # Process body paragraphs
    for paragraph in doc.paragraphs:
        quote_counter = _process_paragraph(paragraph, stats, quote_counter, quote_font)

    # Process tables
    for table in doc.tables:
        quote_counter = _process_table(table, stats, quote_counter, quote_font)

    # Process headers/footers
    if strip_headers_footers:
        for section in doc.sections:
            sectPr = section._sectPr
            for ref in sectPr.findall(qn("w:headerReference")):
                sectPr.remove(ref)
            for ref in sectPr.findall(qn("w:footerReference")):
                sectPr.remove(ref)
    elif process_headers_footers:
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    quote_counter = _process_paragraph(paragraph, stats, quote_counter, quote_font)
                for table in section.header.tables:
                    quote_counter = _process_table(table, stats, quote_counter, quote_font)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    quote_counter = _process_paragraph(paragraph, stats, quote_counter, quote_font)
                for table in section.footer.tables:
                    quote_counter = _process_table(table, stats, quote_counter, quote_font)

    # Save to bytes
    buf = BytesIO()
    doc.save(buf)
    return FormatResult(data=buf.getvalue(), stats=stats)


# == Text extraction ===========================================================

# W namespace for raw XML access
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _wqn(tag: str) -> str:
    """Quick namespace helper for w: prefix."""
    return f"{{{_W_NS}}}{tag}"


def _extract_paragraphs_raw(doc_bytes: bytes) -> list[dict]:
    """Extract paragraphs via raw XML for maximum fidelity."""
    paragraphs = []
    with zipfile.ZipFile(BytesIO(doc_bytes), "r") as zf:
        doc_xml = zf.read("word/document.xml")
        # Build style ID → name map
        styles_map = {}
        if "word/styles.xml" in zf.namelist():
            styles_xml = zf.read("word/styles.xml")
            stree = etree.fromstring(styles_xml)
            for s in stree.iter(_wqn("style")):
                sid = s.get(_wqn("styleId"), "")
                name_elem = s.find(_wqn("name"))
                name = name_elem.get(_wqn("val"), sid) if name_elem is not None else sid
                styles_map[sid] = name

        tree = etree.fromstring(doc_xml)
        body = tree.find(_wqn("body"))
        if body is None:
            return paragraphs

        for para in body.iter(_wqn("p")):
            ppr = para.find(_wqn("pPr"))
            style_id = ""
            outline_level = -1
            if ppr is not None:
                ps = ppr.find(_wqn("pStyle"))
                if ps is not None:
                    style_id = ps.get(_wqn("val"), "")
                ol = ppr.find(_wqn("outlineLvl"))
                if ol is not None:
                    outline_level = int(ol.get(_wqn("val"), "-1"))

            style_name = styles_map.get(style_id, style_id)

            level = -1
            heading_match = re.match(r"[Hh]eading\s*(\d+)", style_name)
            if heading_match:
                level = int(heading_match.group(1))
            elif outline_level >= 0:
                level = outline_level + 1

            texts = []
            for t in para.iter(_wqn("t")):
                if t.text:
                    texts.append(t.text)
            text = "".join(texts).strip()

            if text:
                paragraphs.append({"style": style_name, "text": text, "level": level})

    return paragraphs


def _paragraphs_to_markdown(paragraphs: list[dict]) -> str:
    """Convert paragraph list to Markdown."""
    lines = []
    for p in paragraphs:
        style = p["style"].lower()
        text = p["text"]
        level = p["level"]

        if level >= 1:
            lines.append(f"\n{'#' * level} {text}\n")
        elif "表" in style or "图" in style or "caption" in style:
            lines.append(f"\n**[{text}]**\n")
        elif "题目" in style or "title" in style:
            lines.append(f"\n**{text}**\n")
        else:
            lines.append(f"\n{text}\n")

    return "\n".join(lines).strip() + "\n"


def extract_text(doc_bytes: bytes) -> str:
    """Extract text from a Word document as Markdown.

    Heading styles are converted to ``#`` headings. Table/figure captions
    are bolded. Custom styles (e.g. ZDWP) are detected via outline level.

    Args:
        doc_bytes: Raw bytes of a .docx file.

    Returns:
        Markdown-formatted text.
    """
    paragraphs = _extract_paragraphs_raw(doc_bytes)
    return _paragraphs_to_markdown(paragraphs)


def extract_paragraphs(doc_bytes: bytes) -> list[dict]:
    """Extract structured paragraph data from a Word document.

    Args:
        doc_bytes: Raw bytes of a .docx file.

    Returns:
        List of dicts, each with keys: style, text, level (-1 if not a heading).
    """
    return _extract_paragraphs_raw(doc_bytes)


def extract_chapters(doc_bytes: bytes) -> list[dict]:
    """Extract and split document by top-level headings.

    Args:
        doc_bytes: Raw bytes of a .docx file.

    Returns:
        List of chapter dicts, each with: title, paragraphs, markdown.
    """
    paragraphs = _extract_paragraphs_raw(doc_bytes)
    chapters = []
    current: dict = {"title": "", "paragraphs": []}

    for p in paragraphs:
        if p["level"] == 1:
            if current["paragraphs"]:
                current["markdown"] = _paragraphs_to_markdown(current["paragraphs"])
                chapters.append(current)
            current = {"title": p["text"], "paragraphs": [p]}
        else:
            current["paragraphs"].append(p)

    if current["paragraphs"]:
        current["markdown"] = _paragraphs_to_markdown(current["paragraphs"])
        chapters.append(current)

    return chapters


# == Style cleanup =============================================================


@dataclass
class CleanupResult:
    """Result of style cleanup operation."""

    data: bytes
    log: list[str] = field(default_factory=list)


# Built-in style IDs that should never be deleted
_BUILTIN_KEEP = {"a", "a0"}


def _get_style_map(styles_tree) -> dict:
    """Extract styleId -> {name, type, basedOn, elem} from styles.xml."""
    result = {}
    for s in styles_tree.findall(f".//{_wqn('style')}"):
        sid = s.get(_wqn("styleId"), "")
        stype = s.get(_wqn("type"), "")
        ne = s.find(_wqn("name"))
        name = ne.get(_wqn("val"), sid) if ne is not None else sid
        based = s.find(_wqn("basedOn"))
        base_id = based.get(_wqn("val"), "") if based is not None else ""
        result[sid] = {"name": name, "type": stype, "basedOn": base_id, "elem": s}
    return result


def _get_used_style_ids(doc_tree) -> set:
    """Find all style IDs referenced in document.xml."""
    used = set()
    for tag in ("pStyle", "rStyle", "tblStyle"):
        for elem in doc_tree.iter(_wqn(tag)):
            used.add(elem.get(_wqn("val"), ""))
    used.discard("")
    return used


def _get_needed_ids(style_map: dict, used_ids: set) -> set:
    """Recursively find all basedOn dependencies."""
    needed = set(used_ids)
    queue = list(used_ids)
    while queue:
        sid = queue.pop()
        info = style_map.get(sid)
        if info and info["basedOn"] and info["basedOn"] not in needed:
            needed.add(info["basedOn"])
            queue.append(info["basedOn"])
    return needed


def cleanup_styles(
    doc_bytes: bytes,
    *,
    renames: dict[str, str] | None = None,
    merges: dict[str, str] | None = None,
    delete_unused: bool = True,
) -> CleanupResult:
    """Clean up styles in a Word document.

    Operations (applied in order):
    1. Merge styles: reassign paragraphs from one style to another.
    2. Rename styles: change display names (safe, does not change IDs).
    3. Delete unused: remove style definitions not referenced by any paragraph.

    Args:
        doc_bytes: Raw bytes of a .docx file.
        renames: ``{old_display_name: new_display_name}`` mapping.
        merges: ``{from_style_id: to_style_id}`` mapping.
        delete_unused: Whether to delete unreferenced style definitions.

    Returns:
        CleanupResult with modified document bytes and operation log.
    """
    # Load all parts from the ZIP
    files = {}
    with zipfile.ZipFile(BytesIO(doc_bytes), "r") as zf:
        for info in zf.infolist():
            files[info.filename] = zf.read(info.filename)

    styles_tree = etree.fromstring(files["word/styles.xml"])
    doc_tree = etree.fromstring(files["word/document.xml"])
    style_map = _get_style_map(styles_tree)
    log: list[str] = []

    # 1. Merge styles
    if merges:
        for from_id, to_id in merges.items():
            if from_id not in style_map:
                log.append(f"Merge skipped: source '{from_id}' not found")
                continue
            if to_id not in style_map:
                log.append(f"Merge skipped: target '{to_id}' not found")
                continue
            count = 0
            for ps in doc_tree.iter(_wqn("pStyle")):
                if ps.get(_wqn("val")) == from_id:
                    ps.set(_wqn("val"), to_id)
                    count += 1
            from_name = style_map[from_id]["name"]
            to_name = style_map[to_id]["name"]
            log.append(f"Merged: {from_name} -> {to_name} ({count} paragraphs)")

    # 2. Rename styles
    if renames:
        name_to_id = {info["name"]: sid for sid, info in style_map.items()}
        for old_name, new_name in renames.items():
            sid = name_to_id.get(old_name)
            if not sid:
                log.append(f"Rename skipped: '{old_name}' not found")
                continue
            ne = style_map[sid]["elem"].find(_wqn("name"))
            if ne is not None:
                ne.set(_wqn("val"), new_name)
            log.append(f"Renamed: {old_name} -> {new_name}")

    # 3. Delete unused styles
    if delete_unused:
        used_ids = _get_used_style_ids(doc_tree)
        needed_ids = _get_needed_ids(style_map, used_ids)
        needed_ids |= _BUILTIN_KEEP
        for sid in style_map:
            if sid.startswith("TOC") or sid in ("a", "a0"):
                needed_ids.add(sid)

        deleted = []
        for sid, info in style_map.items():
            if sid not in needed_ids:
                deleted.append((sid, info["name"]))
                for s in styles_tree.findall(f".//{_wqn('style')}"):
                    if s.get(_wqn("styleId")) == sid:
                        s.getparent().remove(s)
                        break

        if deleted:
            log.append(f"Deleted {len(deleted)} unused styles")

    # Save back to ZIP
    files["word/styles.xml"] = etree.tostring(styles_tree, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["word/document.xml"] = etree.tostring(doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in files.items():
            zf.writestr(name, data)

    return CleanupResult(data=buf.getvalue(), log=log)
